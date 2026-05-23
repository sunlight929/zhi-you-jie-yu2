"""
生成新版讲稿（按新 7 页 PPT 结构）
================================
匹配 知忧解郁_答辩展示_v1.0.pptx 的 7 张幻灯片：
  1  封面  → ① 开场（0:30）
  2  项目背景  → ② 项目背景（1:00）
  3  设计思路  → ③ 数据 + 技术合并（1:00）
  4  ▶ 平台演示  → ④ 平台实操（5:00 大头）
  5  创新点  → ⑤ 创新点（0:50）
  6  团队分工  → ⑥ 团队（0:30）
  7  致谢  → ⑦ 致谢（0:25）

总时长 ≈ 8:45（10 分钟内留 1:15 弹性）
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = Path("/Users/hejingqing/Desktop/比赛/计算机设计大赛/讲稿.docx")
TJU_BLUE = RGBColor(0x00, 0x3F, 0x7E)
TJU_GOLD = RGBColor(0xC4, 0x96, 0x1A)
DEEP_GRAY = RGBColor(0x33, 0x33, 0x33)
LIGHT_GRAY = RGBColor(0x77, 0x77, 0x77)


def set_font(run, size=11, bold=False, color=None,
             name="PingFang SC"):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = name
    if color:
        run.font.color.rgb = color
    # 中文字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = rPr.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), name)


def add_title(doc, text, color=TJU_BLUE, size=18):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=True, color=color)


def add_section_header(doc, label, time_str, char_count):
    """段落标题：⑤ 标题（X:XX，约 N 字）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(label)
    set_font(run, size=14, bold=True, color=TJU_BLUE)
    run2 = p.add_run(f"   {time_str}  ·  约 {char_count} 字")
    set_font(run2, size=10, color=LIGHT_GRAY)


def add_para(doc, text, size=11, indent=0, color=DEEP_GRAY,
             bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.6
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    run.italic = italic


def add_separator(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("─" * 40)
    set_font(run, size=8, color=LIGHT_GRAY)


def add_note(doc, text):
    """画面提示（灰色斜体小字）。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("  〔" + text + "〕")
    set_font(run, size=10, color=LIGHT_GRAY)
    run.italic = True


# ============================================================
# 创建文档
# ============================================================
doc = Document()

# 页面边距
for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

# 主标题
add_title(doc, "知忧·解郁 · 答辩演示讲稿", color=TJU_BLUE, size=20)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("配套 PPT：知忧解郁_答辩展示_v1.0.pptx（7 张）")
set_font(run, size=10, color=LIGHT_GRAY)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("总时长目标 8:45（10 分钟内留 1:15 弹性）")
set_font(run, size=10, color=LIGHT_GRAY)

add_separator(doc)


# ============================================================
# ① 封面（Slide 1）
# ============================================================
add_section_header(doc, "①  封面（Slide 1）", "0:30", "100")
add_note(doc, "PPT 第 1 张 · 双校徽 + 项目名 + 双导师")

add_para(doc,
         "各位评委老师好，我们是来自天津大学医学院的参赛队伍。")
add_para(doc,
         "我们的作品名为《知忧·解郁——多源数据驱动的抑郁症风险识别与决策支持平台》，"
         "参加 2026 中国大学生计算机设计大赛大数据应用赛道，生物与医疗大数据小类。")
add_para(doc,
         "团队由两名临床医学专业本科生组成，由医学院张小臣副研究员、"
         "电气自动化与信息工程学院电子信息工程系张淑芳副教授联合指导，"
         "是一项典型的医工交叉作品。")


# ============================================================
# ② 项目背景（Slide 2）
# ============================================================
add_section_header(doc, "②  项目背景（Slide 2）", "1:00", "230")
add_note(doc, "PPT 第 2 张 · 数据卡 + 项目定位")

add_para(doc,
         "抑郁症是全球第二大致残原因。在我国——")
add_para(doc, "• 抑郁症患者总数已达 9,500 万；", indent=0.3)
add_para(doc, "• 大学生抑郁检出率高达 24.7%（Gao L 2020 Sci Rep Meta 分析）；",
         indent=0.3)
add_para(doc, "• 中老年抑郁症状检出率 15-30%（CHARLS 北京大学）；",
         indent=0.3)
add_para(doc, "• 但精神科医师密度仅 3.4 / 10 万，"
              "临床就诊率不足 10%，规范治疗率不足 1%。",
         indent=0.3)

add_para(doc, "由此带来三大结构性挑战——",
         bold=True, color=TJU_BLUE, size=12)
add_para(doc,
         "识别难、资源缺、分布不均。",
         bold=True, color=TJU_BLUE, size=12)

add_para(doc,
         "在临床实习中，我们看到太多患者错过了最佳干预窗口，"
         "看到农村空巢老人因生活自理能力下降而陷入孤独。")
add_para(doc,
         "本作品依托张小臣老师在研横向项目"
         "「面向脑认知的抑郁情绪识别模型开发」的科研积累，"
         "希望把它转化为面向公众与基层卫生机构的实用工具。")


# ============================================================
# ③ 设计思路（Slide 3）
# ============================================================
add_section_header(doc, "③  设计思路（Slide 3）", "1:00", "220")
add_note(doc, "PPT 第 3 张 · 多源数据 + 技术路线")

add_para(doc, "▎ 多源数据生态  ·  共 14,750 条",
         bold=True, color=TJU_BLUE, size=12)
add_para(doc,
         "平台整合四个独立数据源——")
add_para(doc,
         "真实公开数据 5,750 条：美国 NHANES 2017-2018 (5,068) "
         "+ Mendeley 学生 PHQ-9 数据集 (682)，全部为公开学术数据。",
         indent=0.3)
add_para(doc,
         "中国本土流行病学仿真数据 9,000 条：中国大学生 5,000（基于 Gao 2020 Sci Rep）"
         "+ 中老年 4,000（基于 CHARLS 北京大学），参数完全来自已发表流行病学论文。",
         indent=0.3)
add_para(doc,
         "所有数据均使用 PHQ-9 与 CES-D 国际临床金标准量表，统一切点 ≥ 10。",
         indent=0.3, italic=True, color=LIGHT_GRAY)

add_para(doc, "▎ 技术路线  ·  端到端数据闭环",
         bold=True, color=TJU_BLUE, size=12)
add_para(doc,
         "平台采用 Streamlit 多页面 Web 架构，"
         "底层是 scikit-learn 三模型集成——"
         "Random Forest 处理非线性、Gradient Boosting 追求高精度、"
         "Logistic Regression 提供可解释性，"
         "再与 PHQ-9 临床切点按 6 比 4 加权融合，"
         "内部验证集 AUC 达 0.935，召回率 0.866。")


# ============================================================
# ④ 平台实操演示（Slide 4）
# ============================================================
add_section_header(doc, "④  ▶ 进入平台演示（Slide 4 → 实际平台浏览器）",
                   "5:00", "1100")
add_note(doc, "PPT 第 4 张点击中央蓝色按钮 → 跳转 localhost:8501，对着平台讲解")

# Part 4.1
add_para(doc, "【Part 1】 数据驾驶舱（约 50 秒）",
         bold=True, color=TJU_GOLD, size=12)
add_note(doc, "鼠标自上而下扫描大屏")
add_para(doc,
         "这是平台的数据驾驶舱大屏。顶部 5 个 KPI 显示——"
         "总样本量 14,750 条、真实公开 5,750、中国仿真 9,000、"
         "综合检出率 19.5%、模型最佳 AUC 0.935。")
add_para(doc,
         "核心是这张全国 34 省抑郁检出率热力图，使用 Apache ECharts 与"
         "国家测绘地理信息局公开 GeoJSON 渲染。鼠标悬停每个省份显示具体数值，"
         "比如北京市 8.20%、天津市 7.90%。")
add_para(doc,
         "大屏共 9 个数据组件，趋势曲线、PHQ-9 五级分布、"
         "天津 16 区县排行、风险雷达、相关性热力图、性别 × 抑郁堆叠柱，"
         "全部联动呈现，构成完整的数据生态全景。")

add_separator(doc)

# Part 4.2
add_para(doc, "【Part 2】 数据全景（约 50 秒）",
         bold=True, color=TJU_GOLD, size=12)
add_note(doc, "点击「📊 数据全景」→ 切换四个 Tab")
add_para(doc,
         "数据全景页有 4 个 Tab，体现我们对数据透明度的严格态度。")
add_para(doc,
         "第一是真实公开数据：美国 NHANES 5,068 条 PHQ-9 调查 + Mendeley 学生 682 条"
         "（CC BY 4.0 开放许可）。")
add_para(doc,
         "第二、第三是中国学生与中老年仿真数据，参数完全基于已发表论文构造。")
add_para(doc,
         "第四 Tab 是四源对比表。我们在驾驶舱有专门的「方法学说明」面板，"
         "解释跨数据集检出率差异——这是不同人群结构与抽样设计造成的合理结果。")

add_separator(doc)

# Part 4.3
add_para(doc, "【Part 3】 风险因素分析（约 40 秒）",
         bold=True, color=TJU_GOLD, size=12)
add_note(doc, "点击「🔍 风险因素分析」")
add_para(doc,
         "这一页用机器学习量化了风险因子。相关性热力图清晰显示——"
         "社会支持、亲密关系与 PHQ-9 强负相关；"
         "学业压力、经济压力、童年逆境强正相关。")
add_para(doc,
         "三个模型一致认为：学生群体最重要的因子是社会支持、经济压力、睡眠时长；"
         "中老年最重要的是 ADL 自理能力、睡眠时长、慢性病数量——"
         "这与张小臣老师课题组的科研结论高度一致。")

add_separator(doc)

# Part 4.4
add_para(doc, "【Part 4】 智能评估 · 高风险案例 ⭐（约 80 秒）",
         bold=True, color=TJU_GOLD, size=12)
add_note(doc, "点击「🧠 智能评估」→ 填写 PHQ-9 高风险样例 → 提交")
add_para(doc,
         "这是平台最核心的功能：临床三合一评估——"
         "PHQ-9 量表 + 风险因素问卷 + 三模型集成。")
add_para(doc,
         "Case 1：模拟一位大四女生，PHQ-9 几乎全选「几乎每天」。"
         "我们用的是 Kroenke 2001 PHQ-9 量表，这是国家卫健委、"
         "北京安定医院在用的临床金标准。"
         "特别注意第 9 题「自伤念头」，我们做了独立预警机制。")
add_para(doc,
         "提交后平台返回三层结果：PHQ-9 总分 24/27，重度抑郁；"
         "机器学习风险 92%；综合判读：量表与模型一致提示中重度及以上抑郁。")
add_para(doc,
         "因为第 9 题大于 0，平台强制弹出红色紧急提醒，"
         "给出全国 24 小时心理援助热线 400-161-9995 与天津本地热线。",
         color=TJU_GOLD, bold=True)
add_para(doc,
         "贡献分解条形图清晰显示：童年逆境、经济压力、学业压力贡献最大；"
         "社会支持薄弱、睡眠不足进一步加剧风险。")

add_separator(doc)

# Part 4.5
add_para(doc, "【Part 5】 智能评估 · 低风险对照（约 30 秒）",
         bold=True, color=TJU_GOLD, size=12)
add_note(doc, "清空 → 输入低风险样例")
add_para(doc,
         "Case 2：同一份问卷，输入一位生活规律、社会支持充足的大二男生——"
         "PHQ-9 总分 1，模型给出 6.4% 的低风险。")
add_para(doc,
         "这种高低对照体现了平台稳健性——"
         "我们用 PHQ-9 临床切点作为最终判定主导，"
         "确保不会出现「量表为 0 但模型说高风险」的逻辑矛盾。")

add_separator(doc)

# Part 4.6
add_para(doc, "【Part 6】 天津决策支持（约 50 秒）",
         bold=True, color=TJU_GOLD, size=12)
add_note(doc, "点击「🏙️ 天津决策支持」")
add_para(doc,
         "作为天津大学的学生，我们对天津 16 区县做了本地化决策支持。"
         "区县按检出率分四级，每个等级给出具体干预建议——"
         "「资源紧急扩容」、「老年照护」、「转诊网络加强」、「常规健康教育」。")
add_para(doc,
         "散点图清晰显示：精神卫生资源越少、老龄化程度越高的区县，"
         "抑郁检出率越高——印证资源应当向农村和老龄化严重的区倾斜。")
add_para(doc,
         "平台还整合了天津本地求助资源——天津市安定医院、"
         "天津医科大学总医院心理科、第四中心医院心身医学科，"
         "加上 24 小时心理援助热线，形成完整就医闭环。")


# ============================================================
# ⑤ 创新点（Slide 5）
# ============================================================
add_section_header(doc, "⑤  创新点（Slide 5）", "0:50", "200")
add_note(doc, "PPT 第 5 张 · 5 大创新编号卡")

add_para(doc,
         "回到 PPT 总结一下五个创新点——")
add_para(doc, "1. 多源数据生态：真实公开 + 中国仿真双轨，每个数据源透明标注；",
         indent=0.3)
add_para(doc, "2. 临床金标准量表完整集成：PHQ-9 + CES-D 10 原题原选项呈现；",
         indent=0.3)
add_para(doc, "3. 三模型集成 + 量表临床切点融合（6:4 加权），AUC 0.935；",
         indent=0.3)
add_para(doc, "4. 局部贡献可解释性：每位用户都附带「为什么风险高 / 低」的可视化解释；",
         indent=0.3)
add_para(doc, "5. 中国 34 省地图 + 天津 16 区县本地化决策支持。",
         indent=0.3)

add_para(doc,
         "AI 工具协作严格遵守大赛附件 3 的 15 款合规清单——"
         "DeepSeek 加速代码、Kimi 撰写报告、通义千问辅助探索、智谱 AI 优化量表。"
         "所有 AI 输出经队员人工核验，双导师对医学和 AI 内容最终把关。")


# ============================================================
# ⑥ 团队分工（Slide 6）
# ============================================================
add_section_header(doc, "⑥  团队分工（Slide 6）", "0:30", "100")
add_note(doc, "PPT 第 6 张 · 学生 + 双导师")

add_para(doc,
         "团队的医工交叉特色，正是双导师配置赋予的——")
add_para(doc,
         "张小臣副研究员长期研究抑郁症神经机制，正是「面向脑认知的抑郁情绪识别模型」"
         "项目负责人，为我们提供医学专业方向把关；",
         indent=0.3)
add_para(doc,
         "张淑芳副教授是电子信息工程系主任，在人工智能、AIGC 大模型领域积累深厚，"
         "为我们提供 AI 技术方向把关。",
         indent=0.3)


# ============================================================
# ⑦ 致谢（Slide 7）
# ============================================================
add_section_header(doc, "⑦  致谢（Slide 7）", "0:25", "120")
add_note(doc, "PPT 第 7 张 · 校徽 + 张小臣老师引言")

add_para(doc,
         "在临床见习时，张小臣老师说过一句让我印象很深的话——")
add_para(doc,
         "「识别一个抑郁症病人，可能就救一条命。」",
         bold=True, color=TJU_BLUE, size=14, indent=0.5)
add_para(doc,
         "我们希望这个平台能成为更多医生、老师、家人手中的一个小工具，"
         "让被忧愁笼罩的人，早一点被看见。")
add_para(doc,
         "感谢中国大学生计算机设计大赛组委会与南开大学，"
         "感谢张小臣副研究员和张淑芳副教授的指导，"
         "感谢评委老师的审阅。",
         italic=True)
add_para(doc,
         "谢谢评委老师！",
         bold=True, color=TJU_BLUE, size=14)


# ============================================================
# 末尾：录制要点
# ============================================================
add_separator(doc)
add_title(doc, "📋  录制操作要点", color=TJU_BLUE, size=14)

add_para(doc, "▎ 录制策略",
         bold=True, color=TJU_BLUE, size=12)
add_para(doc, "• 第 1-3 段：PPT 全屏播放（出镜 + PPT 切换），约 2:30；",
         indent=0.3)
add_para(doc, "• 第 4 段：点击 PPT 第 4 张的链接 → 切到浏览器 → 屏幕录制讲解，约 5:00；",
         indent=0.3)
add_para(doc, "• 第 5-7 段：切回 PPT 全屏（出镜），约 1:15；",
         indent=0.3)
add_para(doc, "• 总时长 ≈ 8:45。",
         indent=0.3)

add_para(doc, "▎ 准备清单",
         bold=True, color=TJU_BLUE, size=12)
add_para(doc, "• Chrome 用无痕模式打开 localhost:8501，避免插件浮窗；",
         indent=0.3)
add_para(doc, "• 浏览器宽度 1920px，缩放 100%，左侧 Streamlit 侧边栏点 « 折叠；",
         indent=0.3)
add_para(doc, "• 用领夹麦或耳麦，关闭桌面通知和网络（防弹窗）；",
         indent=0.3)
add_para(doc, "• 录制工具：QuickTime（最稳）/ OBS（更专业）；",
         indent=0.3)
add_para(doc, "• 输出 MP4 / H.264 / 1080p，体积 ≤ 200 MB；",
         indent=0.3)
add_para(doc, "• 命名：知忧解郁_演示视频_v1.0.mp4。",
         indent=0.3)

add_para(doc, "▎ 答辩 Q&A 备战",
         bold=True, color=TJU_BLUE, size=12)
add_para(doc, "• Mendeley 数据：是 Mendeley 公开学术数据集（kkzjk253cy，CC BY 4.0），"
              "孟加拉国 IUBAT 团队收集，用于外部验证与跨人群泛化测试。",
         indent=0.3)
add_para(doc, "• 临床医学背景怎么做代码：得益于双导师配置 + 大赛 15 款合规 AI 工具协作，"
              "所有医学专业表达由我们人工核验。",
         indent=0.3)
add_para(doc, "• 检出率差异：不同人群结构与抽样设计的合理结果，"
              "所有数据均使用 PHQ-9 国际标准量表与统一切点。",
         indent=0.3)


# ============================================================
# 保存
# ============================================================
doc.save(str(OUT))
print(f"✅ 讲稿已生成：{OUT}")
print(f"   文件大小：{OUT.stat().st_size // 1024} KB")
